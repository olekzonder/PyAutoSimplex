from docx import Document
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class DOCXGen:
    def __init__(self,filename, data, type, lessThanList):
        self.data = data
        self.filename = filename
        self.type = type
        self.lessThanList = lessThanList
        self.document = Document()
        self.document.add_heading('Табличный симплекс-метод.')
        self.document.add_paragraph('Решим прямую задачу линейного программирования симплексным методом, с использованием симплексной таблицы.')
        s = 'минимальное' if self.type=='min' else 'максимальное'
        f = self.makeFunction(self.data['function'])
        p=self.document.add_paragraph('Определим '+s+' значение целевой функции:') 
        p.add_run('\n F(X)='+f).bold=True
        p.add_run(" при следующих условиях-ограничениях:")
        for i in range(len(data['constraints'])):
            self.document.add_paragraph(self.makeFunction(data['constraints'][i],res=data['resources'][i],lessThan=(1 if i in self.lessThanList else 0)))
            
    def makeFunction(self,f,res = None,lessThan=None,k=None):
        s = ''
        for i in range(len(f)-1):
            if f[i] == 0:
                continue
            x = str(f[i]) if abs(f[i]) != 1.0 else ''
            s +=  x + 'X'+str(i+1) 
            if str(f[i+1])[0] != ' - ':
                s += ' + '
        if f[len(f)-1] != 0:
            s += str(round(f[len(f)-1])) if abs(f[len(f)-1]) != 1 else '' 
            if k == None:
                s += 'X'+str(len(f))
            else:
                s += 'X'+str(k)
        if res == None:
            s += ' -> '
            s += self.type
        if res != None and lessThan != None:
            s += '≥' if lessThan == 0 else '≤'
            s += str(res)
        if res != None and lessThan == None:
            s += '='
            s += str(res)
        return s
        
    def notOptimal(self):
        s = 'положительные' if self.type=='min' else 'отрицательные'
        self.document.add_paragraph("Данный план не является оптимальным, так как индексная строка содержит "+s+" элементы.")
    def optimal(self):
        s = 'положительных' if self.type=='min' else 'отрицательных'
        self.document.add_paragraph("Был найден оптимальный план - индексная строка не содержит "+s+" элементов. Задача линейного программирования решена.")        

    def addTable(self,data):
        simplex = data['data']
        row = -1
        col = -1
        if 'row' in data:
            row = data['row']+1
            col = data['col']+2
        table = self.document.add_table(rows=0, cols=len(simplex[0]))
        for i in range(len(simplex)):
            row_cells = table.add_row().cells
            for j in range(len(simplex[i])):
                row_cells[j].text = str(simplex[i][j])
                if j == col or i == row:
                    xml_elm = row_cells[j]._tc
                    xml_prop = xml_elm.get_or_add_tcPr()
                    shade_obj = OxmlElement('w:shd')
                    shade_obj.set(qn('w:fill'),"99db87")
                    xml_prop.append(shade_obj)
    def canonical(self):
        k  = len(self.data['function'])+1
        sp = []
        if self.type == 'max':
            for i in range(len(self.data['constraints'])):
                print(self.data['constraints'][i])
                if i in self.lessThanList:
                    self.data['constraints'][i].append(1)
                else:
                    self.data['constraints'][i].append(-1)
                self.document.add_paragraph(self.makeFunction(self.data['constraints'][i],res=self.data['resources'][i],k=k))
                sp.append('x'+str(k))
                k+=1
        if self.type == 'min': 
            for i in range(len(self.data['constraints'])):
                print(self.data['constraints'][i])
                if i not in self.lessThanList:
                    self.data['constraints'][i].append(1)
                else:
                    self.data['constraints'][i].append(-1)
                self.document.add_paragraph(self.makeFunction(self.data['constraints'][i],res=self.data['resources'][i],k=k))
                sp.append('x'+str(k))
                k+=1
        sp = ', '.join(sp)
        self.document.add_paragraph("Решим систему уравнений относительно базисных переменных: "+sp + ".")
    def generate(self,data,res):
        for i in range(len(data)):
            if i > 1:
                self.document.add_paragraph("Формируем следующую часть симплексной таблицы.")
                self.document.add_paragraph("Получаем новую симплекс-таблицу:")
            elif i == 1:
                self.document.add_paragraph('Рассмотрим его:')
            else:
                self.document.add_paragraph("Приведём систему неравенств к системе уравнений путём введения дополнительных переменных:")
                self.canonical()
                self.document.add_paragraph("Полагая, что свободные переменные равны 0, получим первый опорный план:")
            self.addTable(data[i])
            if i < len(data)-1 and i > 0:
                self.notOptimal()
                if 'row' in data[i]:
                    row = data[i]['row']+1
                    col = data[i]['col']+2
                    minelm = data[i]['data'][row][len(data[i])-1]
                    s = 'наибольший положительный' if self.type=='min' else 'наименьший отрицательный'
                    self.document.add_paragraph("В качестве ведущего выберем столбец, соответствующий переменной x"+str(col-2)+ ", так как в индексной строке это "+ s +" элемент.")
                    s = '-ая' if col != 3 else '-я'
                    self.document.add_paragraph("Ведущей является " + str(row)+s+" строка.")
                    self.document.add_paragraph("Разрешающий элемент, равный, "+str(data[i]['data'][row][col])+" находится на пересечении ведущего столбца и ведущей строки.")
        self.optimal()
        self.document.add_paragraph("Ответ:")
        for i in range(len(res['result'])):
            p=self.document.add_paragraph("X"+str(i+1)+"=")
            p.add_run(str(res['result'][i])).bold=True
        p = self.document.add_paragraph("F(X)= ")
        p.add_run(str(res['fx'])).bold=True
        self.document.add_paragraph("Решение создано программой PyAutoSimplex (https://github.com/olekzonder/pyautosimplex)")
    def save(self):
        self.document.save(self.filename)